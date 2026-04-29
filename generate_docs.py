

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color_hex):
    
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            element = OxmlElement("w:{}".format(edge))
            for attr_name, attr_val in kwargs[edge].items():
                element.set(qn("w:{}".format(attr_name)), str(attr_val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_page_break(doc):
    
    doc.add_page_break()


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def add_heading_styled(doc, text, level=1):
    
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=(26, 26, 46))
    elif level == 2:
        set_run_font(run, size=14, bold=True, color=(26, 26, 46))
    elif level == 3:
        set_run_font(run, size=12, bold=True, color=(26, 26, 46))
    return heading


def add_paragraph_styled(doc, text, bold=False, italic=False, size=12, alignment=None, space_after=6, first_line_indent=None):
    
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)
    return para


def add_bullet_point(doc, text, bold_prefix=None, level=0):
    
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = para.add_run(bold_prefix)
        set_run_font(run_bold, size=12, bold=True)
        run_rest = para.add_run(text)
        set_run_font(run_rest, size=12)
    else:
        run = para.add_run(text)
        set_run_font(run, size=12)
    if level > 0:
        para.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return para


def add_code_block(doc, code_text):
    
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(code_text)
    set_run_font(run, name="Courier New", size=9, color=(40, 40, 40))
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F0F0F0")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    return para


def create_styled_table(doc, headers, rows, col_widths=None):
    
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(header_text)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A1A2E")

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text))
            set_run_font(run, size=10)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F4F4F8")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement("w:{}".format(border_name))
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        borders.append(border)
    tblPr.append(borders)

    return table


SCREENSHOT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs", "screenshots")

SCREENSHOT_MAP = {
    "5.1.a": ["1. Főoldal — hero.png"],
    "5.1.b": ["1. Főoldal — videók.png"],
    "5.1.c": ["1. Főoldal — terkép.png"],
    "5.2.a": ["2. Regisztráció.png"],
    "5.2.b": ["2. Bejelentkezés.png"],
    "5.3":   ["3. Képek menü.png"],
    "5.4":   ["4. Kapcsolat.png"],
    "5.5":   ["5. Üzenetek táblázat.png"],
    "5.6.a": ["6. CRUD lista.png"],
    "5.6.b": ["5.6.b uresűrlap.png"],
    "5.6.c": ["5.6.cszerkesztes.png"],
    "5.6.d": ["5.6.d torles.png"],
    "7":     ["7. zarthamburger.PNG", "7. nyitotthamb.PNG", "7. képfügg.PNG"],
}


def add_screenshot(doc, key, caption=""):
    files = SCREENSHOT_MAP.get(key, [])
    inserted = False
    for fname in files:
        full = os.path.join(SCREENSHOT_DIR, fname)
        if os.path.exists(full):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run()
            try:
                run.add_picture(full, width=Cm(14))
            except Exception:
                run.add_text(f"[KEPHIBA: {fname}]")
            inserted = True

    if not inserted:
        add_screenshot_placeholder(doc, caption)
        return

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_after = Pt(12)
        run2 = cap_para.add_run(caption)
        set_run_font(run2, size=10, italic=True, color=(100, 100, 100))


def add_screenshot_placeholder(doc, caption=""):

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("[SCREENSHOT PLACEHOLDER]")
    set_run_font(run, size=11, italic=True, color=(150, 150, 150))

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_after = Pt(12)
        run2 = cap_para.add_run(caption)
        set_run_font(run2, size=10, italic=True, color=(100, 100, 100))


def generate_documentation():
    
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rFonts = style.element.rPr.find(qn("w:rFonts")) if style.element.rPr is not None else None
    if rFonts is None:
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")

    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Filmtár")
    set_run_font(run, size=36, bold=True, color=(26, 26, 46))

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Webalkalmazás Dokumentáció")
    set_run_font(run, size=24, bold=True, color=(226, 182, 22))

    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_para.add_run("_" * 50)
    set_run_font(run, size=12, color=(200, 200, 200))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Webprogramozás 1 - Gyakorlat beadandó")
    set_run_font(run, size=16, color=(80, 80, 80))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Készítették:")
    set_run_font(run, size=14, color=(100, 100, 100))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gaál Péter")
    set_run_font(run, size=18, bold=True, color=(26, 26, 46))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Neptun: GULX05")
    set_run_font(run, size=12, color=(100, 100, 100))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Molnár Ádám")
    set_run_font(run, size=18, bold=True, color=(26, 26, 46))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Neptun: MFG82Z")
    set_run_font(run, size=12, color=(100, 100, 100))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026")
    set_run_font(run, size=16, color=(100, 100, 100))

    add_page_break(doc)

    add_heading_styled(doc, "Tartalomjegyzék", level=1)

    toc_items = [
        ("1.", "Projekt áttekintése", "3"),
        ("1.1.", "Az alkalmazás célja és funkciói", "3"),
        ("1.2.", "Témaválasztás indoklása", "3"),
        ("1.3.", "Használt technológiák", "4"),
        ("1.4.", "Front-controller tervezési minta", "4"),
        ("2.", "Rendszerkövetelmények", "5"),
        ("2.1.", "Szerver oldali követelmények", "5"),
        ("2.2.", "Kliens oldali követelmények", "5"),
        ("3.", "Adatbázis tervezés", "6"),
        ("3.1.", "ER diagram", "6"),
        ("3.2.", "Táblák részletes leírása", "6"),
        ("3.3.", "Kapcsolatok (külső kulcsok)", "8"),
        ("4.", "Alkalmazás felépítése", "9"),
        ("4.1.", "Mappastruktúra", "9"),
        ("4.2.", "Fájlok és szerepük", "9"),
        ("4.3.", "Front-controller minta működése", "10"),
        ("4.4.", "PRG (Post-Redirect-Get) minta", "10"),
        ("5.", "Funkciók bemutatása", "11"),
        ("5.1.", "Főoldal", "11"),
        ("5.2.", "Regisztráció és bejelentkezés", "11"),
        ("5.3.", "Képgaléria és feltöltés", "12"),
        ("5.4.", "Kapcsolati űrlap", "13"),
        ("5.5.", "Üzenetek oldal", "14"),
        ("5.6.", "CRUD műveletek", "14"),
        ("6.", "Biztonság", "16"),
        ("7.", "Reszponzív dizájn", "17"),
        ("8.", "Összefoglalás", "18"),
        ("9.", "Munkafelosztás", "19"),
        ("9.1.", "Gaál Péter felelősségi területe", "19"),
        ("9.2.", "Molnár Ádám felelősségi területe", "19"),
        ("9.3.", "Közös munka", "20"),
        ("10.", "Beadási és belépési adatok", "21"),
        ("10.1.", "Internetes elérhetőség", "21"),
        ("10.2.", "Tárhely (FTP) belépési adatok", "21"),
        ("10.3.", "MySQL adatbázis belépési adatok", "21"),
        ("10.4.", "Teszt-felhasználók", "22"),
        ("11.", "Irodalomjegyzék", "23"),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_num = p.add_run(num + " ")
        is_main = not "." in num[:-1]  # Fő fejezet-e
        set_run_font(run_num, size=12, bold=is_main)
        run_title = p.add_run(title)
        set_run_font(run_title, size=12, bold=is_main)
        dots = " " + "." * (60 - len(num) - len(title)) + " "
        run_dots = p.add_run(dots)
        set_run_font(run_dots, size=12, color=(180, 180, 180))
        run_page = p.add_run(page)
        set_run_font(run_page, size=12)
        if num.count(".") > 1:
            p.paragraph_format.left_indent = Cm(1)

    add_page_break(doc)

    add_heading_styled(doc, "1. Projekt áttekintése", level=1)

    add_heading_styled(doc, "1.1. Az alkalmazás célja és funkciói", level=2)

    add_paragraph_styled(doc,
        "A Filmtár egy teljes értékű webalkalmazás, amely egy filmadatbázis "
        "kezelését teszi lehetővé. Az alkalmazás célja, hogy a felhasználók "
        "kényelmesen tallózhatnak a magyar filmek között, uj filmeket "
        "vehetnek fel az adatbázisba, szerkeszthetik és torolhetik azokat. "
        "Az alkalmazás emellett közösségi funkciókal is rendelkezik: "
        "felhasználói regisztráció, bejelentkezés, képfeltöltés és "
        "kapcsolati űrlap reven.")

    add_paragraph_styled(doc, "Az alkalmazás fő funkciói:")

    features = [
        "Felhasználói regisztráció és bejelentkezés (session alapú autentikáció)",
        "Filmek teljes koru kezelese (CRUD: létrehozas, olvasás, módosítás, törlés)",
        "Képgaléria képfeltöltési lehetőséggel (csak bejelentkezett felhasználóknak)",
        "Kapcsolati űrlap ketszintu (kliens + szerver oldali) validációval",
        "Üzenetek megtekintese (bejelentkezett felhasználóknak)",
        "Főoldal hero szekcióval, videó beágyazással és Google Maps térképpel",
        "Reszponzív (mobilbarat) megjelenes hamburger menüvel",
    ]
    for f in features:
        add_bullet_point(doc, f)

    add_heading_styled(doc, "1.2. Témaválasztás indoklása", level=2)

    add_paragraph_styled(doc,
        "A Filmtár témaválasztás mellett több szempont is szolt. A filmadatbázis "
        "egy kozelitheto és koennyen ertelmezheto téma, amely lehetőséget ad "
        "a CRUD műveletek termeszetes bemutatasa. A filmeknek jol definialt "
        "tulajdonsagai vannak (cim, rendező, ev, műfaj, értékelés, leírás), "
        "amelyek kivaloan illusztraljak az adatbázis-tervezés és az "
        "űrlapkezeles elveit.")

    add_paragraph_styled(doc,
        "A magyar filmek választása kulonleges jelleget ad a projektnek, "
        "és egyben kulturalis értéket is kozvetit. Az adatbázisban 12 ismert "
        "magyar film szerepel, koztuk Oscar-dijas alkotasok (Saul fia, Mindenki) "
        "és ikonikus klasszikusok (A tanu, Macskafogó).")

    add_heading_styled(doc, "1.3. Használt technológiák", level=2)

    tech_data = [
        ["PHP 8+", "Szerver oldali programozasi nyelv", "Üzleti logika, adatbázis műveletek, session kezeles"],
        ["MySQL 5.7+", "Relacios adatbázis-kezelő rendszer", "Adattárolas (felhasználók, filmek, üzenetek, képek)"],
        ["HTML5", "Jelolonyelv", "Oldal strukturaja, szemantikus elemek"],
        ["CSS3", "Stíluslapnyelv", "Megjelenes, reszponzív design, animációk"],
        ["JavaScript", "Kliens oldali programozasi nyelv", "Form validáció, hamburger menü, lightbox galéria"],
        ["PDO", "PHP Data Objects", "Adatbázis absztrakcios reteg (prepared statements)"],
    ]

    create_styled_table(doc,
        ["Technologia", "Típus", "Felhasználasi terület"],
        tech_data,
        col_widths=[4, 5, 8])

    doc.add_paragraph()

    add_heading_styled(doc, "1.4. Front-controller tervezési minta", level=2)

    add_paragraph_styled(doc,
        "Az alkalmazás a front-controller tervezési mintát valósítja meg. "
        "Ez azt jelenti, hogy minden HTTP kérés egyetlen belépési ponton, "
        "az index.php fájlon keresztül erkezik. Az index.php a $_GET['page'] "
        "paraméter alapján donti el, mely oldal tartalom jelenjen meg. "
        "Ez a megoldasnak szamos elonye van:")

    fc_advantages = [
        "Központi hozzaferes-kezelest biztosít (pl. session ellenőrzés)",
        "A POST kérésesek feldolgozása egységes helyen történik (PRG minta)",
        "A navigáció és a sablonok (header/footer) automatikusan betoltodnek",
        "Biztonsági ellenőrzés: csak az engedélyezett oldalak érhetők el",
        "Konnyebb karbantartas: egyetlen fájlban látszik a teljes routing logika",
    ]
    for adv in fc_advantages:
        add_bullet_point(doc, adv)

    add_page_break(doc)

    add_heading_styled(doc, "2. Rendszerkövetelmények", level=1)

    add_heading_styled(doc, "2.1. Szerver oldali követelmények", level=2)

    server_reqs = [
        ["Apache", "2.4+", "Webszerver, mod_rewrite tamogatassal"],
        ["PHP", "8.0+", "Szerver oldali szkriptnyelv"],
        ["MySQL", "5.7+ / MariaDB 10.3+", "Relacios adatbázis-kezelő"],
        ["PDO MySQL", "Beepitett", "PHP adatbázis driver (php-mysql csomag)"],
    ]

    create_styled_table(doc,
        ["Szoftver", "Minimum verzio", "Megjegyzés"],
        server_reqs,
        col_widths=[4, 4.5, 8.5])

    doc.add_paragraph()

    add_paragraph_styled(doc,
        "Ajánlott fejlesztoi kornyezet: XAMPP, WAMP, MAMP vagy Laragon, "
        "amelyek az összes szükséges szoftvert tartalmazzak.")

    add_heading_styled(doc, "2.2. Kliens oldali követelmények", level=2)

    client_reqs = [
        ["Google Chrome", "90+", "Teljes tamogatas"],
        ["Mozilla Firefox", "88+", "Teljes tamogatas"],
        ["Microsoft Edge", "90+", "Teljes tamogatas (Chromium alapú)"],
        ["Safari", "14+", "Teljes tamogatas"],
    ]

    create_styled_table(doc,
        ["Bongeszo", "Minimum verzio", "Tamogatottsag"],
        client_reqs,
        col_widths=[5, 4, 8])

    doc.add_paragraph()

    add_paragraph_styled(doc,
        "Az alkalmazás modern CSS3 technikat (CSS Grid, Flexbox, CSS változók, "
        "média queries) és HTML5 elemeket használ, ezért modern bongeszo "
        "szükséges a helyes megjeleniteshtez. JavaScript engedélyezése szükséges "
        "a kliens oldali űrlapvalidációhoz, a hamburger menühoz és a lightbox "
        "galériahoz.")

    add_page_break(doc)

    add_heading_styled(doc, "3. Adatbázis tervezés", level=1)

    add_paragraph_styled(doc,
        "Az alkalmazás egy 'filmtar' nevu MySQL adatbázist használ, amely négy "
        "táblát tartalmaz. Az adatbázis utf8mb4 karakterkészlettel és "
        "utf8mb4_hungarian_ci rendezéssel működik, biztosítva a magyar "
        "karakterek helyes kezelését.")

    add_heading_styled(doc, "3.1. ER diagram (szöveges ábrazolas)", level=2)

    er_text = (
        "+------------------+          +------------------+\n"
        "|   felhasználók   |          |      filmek      |\n"
        "+------------------+          +------------------+\n"
        "| PK id            |          | PK id            |\n"
        "|    felhasználónév |          |    cím           |\n"
        "|    jelszó         |          |    rendező       |\n"
        "|    családi_nev    |          |    ev            |\n"
        "|    utónév         |          |    műfaj         |\n"
        "|    email          |          |    értékelés     |\n"
        "|    létrehozva     |          |    leírás        |\n"
        "+--------+---------+          +------------------+\n"
        "         |\n"
        "         | 1:N                       1:N\n"
        "         +-------------------+-------------------+\n"
        "         |                                       |\n"
        "         v                                       v\n"
        "+------------------+          +------------------+\n"
        "|    üzenetek      |          |      képek       |\n"
        "+------------------+          +------------------+\n"
        "| PK id            |          | PK id            |\n"
        "|    nev            |          |    fájlnév       |\n"
        "|    email          |          |    eredeti_nev   |\n"
        "|    tárgy          |          | FK feltöltő_id   |\n"
        "|    üzenet         |          |    feltoltve     |\n"
        "| FK küldő_id      |          +------------------+\n"
        "|    küldve         |\n"
        "+------------------+\n"
    )
    add_code_block(doc, er_text)

    add_paragraph_styled(doc,
        "A felhasználók tábla a központi entitás, amelyhez az üzenetek "
        "és a képek táblák kapcsolódnak külső kulcsokon keresztül. A filmek "
        "tábla önálló entitás, nem kapcsolódik más táblákhoz.")

    add_heading_styled(doc, "3.2. Táblák részletes leírása", level=2)

    add_heading_styled(doc, "felhasználók tábla", level=3)
    add_paragraph_styled(doc,
        "A rendszer regisztralt felhasználóit tárolja. A jelszó mező "
        "PHP password_hash() függvény által generalt bcrypt hash-t tartalmaz.")

    felh_rows = [
        ["id", "INT", "PK, AUTO_INCREMENT", "Egyedi azonosító"],
        ["felhasználónév", "VARCHAR(50)", "NOT NULL, UNIQUE", "Bejelentkezési nev"],
        ["jelszó", "VARCHAR(255)", "NOT NULL", "Bcrypt hash (password_hash)"],
        ["családi_nev", "VARCHAR(100)", "NOT NULL", "Családi név"],
        ["utónév", "VARCHAR(100)", "NOT NULL", "Utónév"],
        ["email", "VARCHAR(100)", "NOT NULL", "E-mail cim"],
        ["létrehozva", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Regisztráció időpontja"],
    ]
    create_styled_table(doc,
        ["Oszlop", "Típus", "Megszorítas", "Leírás"],
        felh_rows,
        col_widths=[3.5, 3.5, 5, 5])

    doc.add_paragraph()

    add_heading_styled(doc, "filmek tábla", level=3)
    add_paragraph_styled(doc,
        "A filmadatbázis központi táblája. Minden film adatait itt tároljuk. "
        "Ez a fő CRUD tábla, amelyen a létrehozas, olvasás, módosítás és "
        "törlés műveletek történnek.")

    filmek_rows = [
        ["id", "INT", "PK, AUTO_INCREMENT", "Egyedi azonosító"],
        ["cim", "VARCHAR(200)", "NOT NULL", "Film cime"],
        ["rendező", "VARCHAR(100)", "NOT NULL", "Rendező neve"],
        ["ev", "INT", "NOT NULL", "Megjelenes eve"],
        ["műfaj", "VARCHAR(100)", "NOT NULL", "Műfaj megnevezése"],
        ["értékelés", "DECIMAL(3,1)", "DEFAULT NULL", "Értékelés (0.0-10.0)"],
        ["leírás", "TEXT", "DEFAULT NULL", "Film rövid leírása"],
    ]
    create_styled_table(doc,
        ["Oszlop", "Típus", "Megszorítas", "Leírás"],
        filmek_rows,
        col_widths=[3.5, 3.5, 5, 5])

    doc.add_paragraph()

    add_heading_styled(doc, "üzenetek tábla", level=3)
    add_paragraph_styled(doc,
        "A kapcsolatfelveteli űrlapon keresztül küldőtt üzeneteket tárolja. "
        "A küldő_id mező NULL értéket vesz fel, ha vendég (nem bejelentkezett "
        "felhasználó) küldte az üzenetet.")

    üzenetek_rows = [
        ["id", "INT", "PK, AUTO_INCREMENT", "Egyedi azonosító"],
        ["nev", "VARCHAR(100)", "NOT NULL", "Küldő neve"],
        ["email", "VARCHAR(100)", "NOT NULL", "Küldő e-mail cime"],
        ["tárgy", "VARCHAR(200)", "NOT NULL", "Üzenet tárgya"],
        ["üzenet", "TEXT", "NOT NULL", "Üzenet szövege"],
        ["küldő_id", "INT", "DEFAULT NULL, FK", "Küldő felhasználó ID-ja"],
        ["küldve", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Küldés időpontja"],
    ]
    create_styled_table(doc,
        ["Oszlop", "Típus", "Megszorítas", "Leírás"],
        üzenetek_rows,
        col_widths=[3.5, 3.5, 5, 5])

    doc.add_paragraph()

    add_heading_styled(doc, "képek tábla", level=3)
    add_paragraph_styled(doc,
        "A felhasználók által feltöltött galéria képek nyilvantartasa. "
        "Csak bejelentkezett felhasználók tölthetnek fel képet, ezért a "
        "feltöltő_id mező kötelező (NOT NULL).")

    képek_rows = [
        ["id", "INT", "PK, AUTO_INCREMENT", "Egyedi azonosító"],
        ["fájlnév", "VARCHAR(255)", "NOT NULL", "Tarolt fájlnév (szerveren)"],
        ["eredeti_nev", "VARCHAR(255)", "NOT NULL", "Eredeti fájlnév (feltöltéskor)"],
        ["feltöltő_id", "INT", "NOT NULL, FK", "Feltöltő felhasználó ID-ja"],
        ["feltoltve", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Feltöltés időpontja"],
    ]
    create_styled_table(doc,
        ["Oszlop", "Típus", "Megszorítas", "Leírás"],
        képek_rows,
        col_widths=[3.5, 3.5, 5, 5])

    add_heading_styled(doc, "3.3. Kapcsolatok (külső kulcsok)", level=2)

    add_paragraph_styled(doc, "Az adatbázisban két külső kulcs (Foreign Key) kapcsolat található:")

    fk_rows = [
        ["üzenetek.küldő_id", "felhasználók.id", "ON DELETE SET NULL", "Felhasználó törlése esetén az üzenet megmarad, de a küldő_id NULL-ra valt"],
        ["képek.feltöltő_id", "felhasználók.id", "ON DELETE CASCADE", "Felhasználó törlése esetén a felhasználó által feltöltött képek is torlodnek"],
    ]
    create_styled_table(doc,
        ["Külső kulcs", "Hivatközött mezo", "Törlési szabály", "Magyarazat"],
        fk_rows,
        col_widths=[3.5, 3.5, 3.5, 6.5])

    doc.add_paragraph()

    add_paragraph_styled(doc,
        "Mindkét kapcsolat ON UPDATE CASCADE szabályt alkalmaz, ami azt jelenti, "
        "hogy ha a felhasználók táblában egy id érték módosul, a kapcsolódó "
        "táblák automatikusan frissülnek.")

    add_page_break(doc)

    add_heading_styled(doc, "4. Alkalmazás felépítése", level=1)

    add_heading_styled(doc, "4.1. Mappastruktúra", level=2)

    add_paragraph_styled(doc,
        "Az alkalmazás következő mappastrukturaba szervezodo:")

    tree = (
        "WebProg/\n"
        "|-- index.php              # Front controller (fő vezerlo)\n"
        "|-- config.php             # Adatbázis konfiguráció (PDO)\n"
        "|\n"
        "|-- css/\n"
        "|   |-- style.css          # Fő stíluslap (reszponzív)\n"
        "|\n"
        "|-- js/\n"
        "|   |-- main.js            # Hamburger menü, flash, lightbox\n"
        "|   |-- validation.js      # Kapcsolati űrlap validáció\n"
        "|\n"
        "|-- pages/\n"
        "|   |-- főoldal.php        # Főoldal (hero, videók, térkép)\n"
        "|   |-- belépés.php        # Bejelentkezési űrlap\n"
        "|   |-- regisztráció.php   # Regisztrációs űrlap\n"
        "|   |-- kijelentkezés.php  # Kijelentkezési oldal\n"
        "|   |-- képek.php          # Képgaléria + feltöltés\n"
        "|   |-- kapcsolat.php      # Kapcsolati űrlap\n"
        "|   |-- üzenetek.php       # Üzenetek listázása\n"
        "|   |-- crud.php           # Filmek CRUD kezelese\n"
        "|\n"
        "|-- templates/\n"
        "|   |-- header.php         # Fejlec sablon (navigáció)\n"
        "|   |-- footer.php         # Lablec sablon\n"
        "|\n"
        "|-- sql/\n"
        "|   |-- database.sql       # Adatbázis létrehozo szkript\n"
        "|   |-- seed.php           # Mintaadatok beszurasa PHP-bol\n"
        "|\n"
        "|-- uploads/               # Feltöltőtt képek könyvtára\n"
        "|-- videós/                # Videó fájlok könyvtára\n"
    )
    add_code_block(doc, tree)

    add_heading_styled(doc, "4.2. Fájlok és szerepük", level=2)

    files_data = [
        ["index.php", "Front controller - egyetlen belépési pont, routing, POST feldolgozás"],
        ["config.php", "PDO adatbázis kapcsolat létrehozása, hibakezeles beállítása"],
        ["css/style.css", "Teljes stíluslap: layout, szinek, reszponzív média queries"],
        ["js/main.js", "Hamburger menü, flash üzenet animáció, lightbox galéria"],
        ["js/validation.js", "Kapcsolati űrlap kliens oldali validációja (JS)"],
        ["pages/főoldal.php", "Főoldal: hero szekci, videók, Google Maps"],
        ["pages/belépés.php", "Bejelentkezési űrlap megjelenítése"],
        ["pages/regisztráció.php", "Regisztrációs űrlap (6 mezo)"],
        ["pages/kijelentkezés.php", "Kijelentkezési visszajelzes oldal"],
        ["pages/képek.php", "Képgaléria + feltöltési űrlap"],
        ["pages/kapcsolat.php", "Kapcsolati űrlap (szerver+kliens validáció)"],
        ["pages/üzenetek.php", "Üzenetek listázása táblázatban"],
        ["pages/crud.php", "Filmek CRUD: listázás, létrehozas, szerkesztés, törlés"],
        ["templates/header.php", "HTML fejlec, navigáció, felhasználói informacio"],
        ["templates/footer.php", "HTML lablec, JS betoltese"],
        ["sql/database.sql", "Adatbázis és táblák létrehozása, mintaadatok (SQL)"],
        ["sql/seed.php", "Adatbázis inicializalas PHP-bol (password_hash)"],
    ]
    create_styled_table(doc,
        ["Fájl", "Szerep / Leírás"],
        files_data,
        col_widths=[5, 12])

    doc.add_paragraph()

    add_heading_styled(doc, "4.3. Front-controller minta működése", level=2)

    add_paragraph_styled(doc,
        "Az index.php az alkalmazás egyetlen belépési pontja. Minden HTTP "
        "kérést ez a fájl fogad és dolgoz fel. A működése a következő:")

    fc_steps = [
        "1. Munkamenet (session) inditasa: session_start()",
        "2. Konfiguráció betoltese: config.php (adatbázis kapcsolat)",
        "3. Segefüggvények definialasa: flash(), getFlash(), bejelentkezveVan(), redirect()",
        "4. POST kérések feldolgozása (ha van $_POST['action']):",
        "   - login: Bejelentkezés feldolgozás",
        "   - register: Regisztráció feldolgozás",
        "   - contact_submit: Kapcsolati űrlap feldolgozás",
        "   - crud_create / crud_update / crud_delete: Film CRUD műveletek",
        "   - image_upload: Képfeltöltés feldolgozás",
        "5. Útvonalválasztás (routing): $_GET['page'] alapján",
        "6. Engedélyezett oldalak ellenőrzése (whitelist)",
        "7. Header sablon betoltese (templates/header.php)",
        "8. Flash üzenet megjelenítése (ha van)",
        "9. Oldal tartalom betoltese (pages/*.php)",
        "10. Footer sablon betoltese (templates/footer.php)",
    ]
    for step in fc_steps:
        if step.startswith("   "):
            add_paragraph_styled(doc, step, size=11, space_after=2)
        else:
            add_bullet_point(doc, step)

    add_heading_styled(doc, "4.4. PRG (Post-Redirect-Get) minta", level=2)

    add_paragraph_styled(doc,
        "Az alkalmazás következetesen alkalmazza a PRG (Post-Redirect-Get) "
        "mintát. Ennek lenyege, hogy a POST kérések feldolgozása után az "
        "alkalmazás HTTP 302 átirányítast hajt vegre. Ez megakadályozza, "
        "hogy az oldal ujratoltesevel (F5) a felhasználó vetelenul ujra "
        "elküldje az űrlapadatokat.")

    add_paragraph_styled(doc, "A PRG minta folyamata:")

    prg_steps = [
        "1. A felhasználó kitolti az űrlapot és megnyomja a Küldés gombot (POST keres)",
        "2. Az index.php feldolgozza a POST adatokat (pl. adatbázisba ment)",
        "3. Flash üzenetet tárol a session-ben (siker vagy hiba)",
        "4. HTTP 302 átirányítast hajt vegre a redirect() függvénnyel",
        "5. A bongeszo automatikusan GET kérést küld az új URL-re",
        "6. Az oldal betoltodik, a flash üzenet megjelenik, majd torlodik a session-bol",
    ]
    for step in prg_steps:
        add_bullet_point(doc, step)

    add_page_break(doc)

    add_heading_styled(doc, "5. Funkciók bemutatása", level=1)

    add_heading_styled(doc, "5.1. Főoldal", level=2)

    add_paragraph_styled(doc,
        "A főoldal (főoldal.php) az alkalmazás nyitóoldala, amely három fő "
        "szekciót tartalmaz:")

    add_heading_styled(doc, "Hero szekció", level=3)
    add_paragraph_styled(doc,
        "Az oldal tetején egy látványos hero szekci fogadja a látogatót. "
        "A 'Filmtár - A kedvenc filmjeid egy helyen' címsort egy bemutató "
        "szöveg követi, amely röviden ismerteti az alkalmazást. Ez alatt "
        "egy 'Rólunk' szekci reszletesebben bemutatja a projekt céljat "
        "és lehetőségeit.")

    add_screenshot(doc, "5.1.a", "5.1.a. ábra: Főoldal - Hero szekció és Rólunk szöveg")

    add_heading_styled(doc, "Videók szekció", level=3)
    add_paragraph_styled(doc,
        "A videók szekció két videót tartalmaz egymás mellett: egy helyi "
        "videólejátszót (<videó> elem, sample.mp4) és egy YouTube "
        "beagyazott videót (<iframe>). Mindkéttő reszponzívan alkalmazkodik "
        "a képernyőmerthez.")

    add_screenshot(doc, "5.1.b", "5.1.b. ábra: Főoldal - Videók szekció")

    add_heading_styled(doc, "Google Maps szekció", level=3)
    add_paragraph_styled(doc,
        "Az oldal aljan egy Google Maps térkép beágyazás lathatoott, amely "
        "egy budapesti cimet (1052 Budapest, Vaci utca 1.) jeleníti meg. "
        "A térkép iframe-ként van beagyazva, lazy loading tamogatassal.")

    add_screenshot(doc, "5.1.c", "5.1.c. ábra: Főoldal - Google Maps beágyazás")

    add_heading_styled(doc, "5.2. Regisztráció és bejelentkezés", level=2)

    add_heading_styled(doc, "Regisztráció", level=3)
    add_paragraph_styled(doc,
        "A regisztrációs oldal (regisztráció.php) egy űrlapot jeleníti meg "
        "hat mezeovel: családi név, utónév, felhasználónév, e-mail, "
        "jelszó és jelszó megerosites. A validáció szerver oldalon történik "
        "az index.php-ban:")

    reg_validations = [
        "Minden mező kitöltése kötelező",
        "E-mail cím formatum ellenőrzése (filter_var, FILTER_VALIDATE_EMAIL)",
        "Jelszavak egyezosenek vizsgalata",
        "Jelszó minimum hossz: 6 karakter",
        "Felhasználónév egyedisegenek ellenőrzése az adatbázisban",
        "Jelszó hashelese: password_hash($jelszó, PASSWORD_DEFAULT)",
    ]
    for v in reg_validations:
        add_bullet_point(doc, v)

    add_screenshot(doc, "5.2.a", "5.2.a. ábra: Regisztrációs űrlap")

    add_heading_styled(doc, "Bejelentkezés", level=3)
    add_paragraph_styled(doc,
        "A bejelentkezési oldal (belépés.php) két mezobol all: felhasználónév "
        "és jelszó. A feldolgozás során a rendszer a felhasználónév alapján "
        "keresi ki a felhasználót az adatbázisbol (prepared statement), majd "
        "a password_verify() függvénnyel ellenőrzi a jelszót. Sikérés "
        "bejelentkezés után a felhasználó adatai a $_SESSION['user'] tömbben "
        "tárolódnak (jelszó nélkül!).")

    add_paragraph_styled(doc,
        "Ha a felhasználó mar be van jelentkezve, az oldal errol tájékoztatja "
        "és nem jelenítit meg az űrlapot.")

    add_screenshot(doc, "5.2.b", "5.2.b. ábra: Bejelentkezési űrlap")

    add_heading_styled(doc, "Kijelentkezés", level=3)
    add_paragraph_styled(doc,
        "A kijelentkezés (kijelentkezés) nem kulon oldal, hanem az index.php "
        "kezeli. A session törlése ($_SESSION = []; session_destroy();) után "
        "uj session indul a flash üzenethez, majd átirányítas történik a "
        "főoldalra. A kijelentkezés.php csak egy visszajelzo oldal, amelyre "
        "normalis esetben nem jutunk el (az átirányítas korabban megtörténik).")

    add_heading_styled(doc, "5.3. Képgaléria és feltöltés", level=2)

    add_paragraph_styled(doc,
        "A képgaléria oldal (képek.php) két fő funkciót lat el: "
        "a feltöltött képek megjelenítéset galéria nézetben és "
        "uj képek feltöltéseteteaet.")

    add_heading_styled(doc, "Galéria nézet", level=3)
    add_paragraph_styled(doc,
        "A képek CSS Grid elrendezésben jelennek meg (3 oszlop asztalon, "
        "2 oszlop tableten, 1 oszlop mobilon). Minden kép alatt látszik "
        "a feltöltő neve és a feltöltés dátuma. A képekre kattintva egy "
        "lightbox overlay nyilik meg a nagyitott keptel (JavaScript által "
        "vezerelt). Az ESC billentyuvel vagy az overlay-re kattintassal "
        "bezarhato.")

    add_heading_styled(doc, "Képfeltöltés", level=3)
    add_paragraph_styled(doc,
        "Csak bejelentkezett felhasználók tölthetnek fel képet. "
        "A feltöltési űrlap enctype='multipart/form-data' attributummal "
        "rendelkezik. A feltöltés során a következő ellenőrzések történnek:")

    upload_checks = [
        "Bejelentkezés ellenőrzése (session)",
        "Fájl erkezesenek ellenőrzése ($_FILES['kep'])",
        "MIME típus ellenőrzése (mime_content_type): image/jpeg, image/png, image/gif, image/webp",
        "Kiterjesztes ellenőrzése: jpg, jpeg, png, gif, webp",
        "Egyedi fájlnév generalasa (uniqid) az ütközések elkerülésére",
        "Fájl athelyezese az uploads/ könyvtárba (move_uploaded_file)",
        "Adatbázisba mentés (képek tábla)",
    ]
    for c in upload_checks:
        add_bullet_point(doc, c)

    add_screenshot(doc, "5.3", "5.3. ábra: Képgaléria CSS Grid elrendezésben")

    add_page_break(doc)

    add_heading_styled(doc, "5.4. Kapcsolati űrlap", level=2)

    add_paragraph_styled(doc,
        "A kapcsolati oldal (kapcsolat.php) egy űrlapot tartalmaz négy "
        "mezővel: nev, e-mail, tárgy és üzenet. Az űrlap kettoa validáción "
        "esik at: kliens oldali (JavaScript) és szerver oldali (PHP).")

    add_heading_styled(doc, "Kliens oldali validáció (validation.js)", level=3)
    add_paragraph_styled(doc,
        "A JavaScript validáció az űrlap submit esemenyre reagal. "
        "A validateContactForm() függvény ellenőrzi az összes mezot, és "
        "hibaüzenetet jelenita meg a mező mellett (span.error elemben). "
        "Nem használ HTML5 validációs attributumokat (required, pattern), "
        "a validáció teljes egeszeben JavaScript-ben történik.")

    js_validations = [
        "Nev: ne legyen ures, minimum 2 karakter",
        "E-mail: ne legyen ures, regex minta: /^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$/",
        "Tárgy: ne legyen ures, minimum 3 karakter",
        "Üzenet: ne legyen ures, minimum 10 karakter",
    ]
    for jv in js_validations:
        add_bullet_point(doc, jv)

    add_heading_styled(doc, "Szerver oldali validáció (index.php)", level=3)
    add_paragraph_styled(doc,
        "A szerver oldali validáció a POST kérés feldolgozásakor fut le. "
        "Ellenorzi a kötelezőe mezőkat és az e-mail formatumot "
        "(filter_var, FILTER_VALIDATE_EMAIL). Hiba esetén a hibaüzenetek "
        "és az űrlapadatok a session-ben tárolódnak, majd átirányítas "
        "történik. Az adatok megorzese lehetővé teszi, hogy a felhasználó "
        "ne veszitse el a mar kitöltött mezők tartalmat.")

    add_paragraph_styled(doc,
        "Bejelentkezett felhasználó esetén a nev és az e-mail mező "
        "automatikusan kitoltodik a session adataival.")

    add_screenshot(doc, "5.4", "5.4. ábra: Kapcsolati űrlap validációs hibaüzenetekkel")

    add_heading_styled(doc, "5.5. Üzenetek oldal", level=2)

    add_paragraph_styled(doc,
        "Az üzenetek oldal (üzenetek.php) a kapcsolati űrlapon küldőtt "
        "üzenetek listajat jelenítia meg táblázatos formatumban. Csak "
        "bejelentkezett felhasználók szamara elérhető - nem bejelentkezett "
        "felhasználók hibaüzenet kapnak.")

    add_paragraph_styled(doc,
        "A táblázat oszlopai: sorszam, küldő neve, e-mail, tárgy, üzenet "
        "szövege és a küldés ideje. A vendég (nem bejelentkezett) küldőnel "
        "'Vendég' felirat jelenik meg. Az üzenetek a küldés dátuma szerint "
        "csokkenosorrendben jelennek meg (legujabb elol).")

    add_screenshot(doc, "5.5", "5.5. ábra: Üzenetek oldal táblázatos nézetben")

    add_heading_styled(doc, "5.6. CRUD műveletek (filmek kezelése)", level=2)

    add_paragraph_styled(doc,
        "A CRUD (Create, Read, Update, Delete) funkciónalitas a filmek "
        "kezelesere szolgal. Egyetlen fájl (crud.php) kezeli mind a négy "
        "műveletet a $_GET['action'] paraméter alapján.")

    add_heading_styled(doc, "Listazas (Read)", level=3)
    add_paragraph_styled(doc,
        "Az alapértelmezett nézet (action=list) az összes filmet egy "
        "táblázatban jelenítia meg. A táblázat oszlopai: #, Cím, Rendező, "
        "Ev, Műfaj, Értékelés, Muveletek. Minden sorban Szerkesztés és "
        "Törlés gombok találhatóok. A filmek id szerint csokkeno sorrendben "
        "jelennek meg (legujabb elol). Az oldal tetején egy 'Uj film "
        "hozzáadása' gomb található.")

    add_screenshot(doc, "5.6.a", "5.6.a. ábra: Filmek listázása táblázatban")

    add_heading_styled(doc, "Létrehozas (Create)", level=3)
    add_paragraph_styled(doc,
        "Az 'Uj film hozzáadása' gombra kattintva (action=uj) egy űrlap "
        "jelenik meg hat mezővel: cim, rendező, ev, műfaj, értékelés és "
        "leírás. A cim, rendező és ev kitöltése kötelező. Az ev mező "
        "number típusu, 1888 és az aktualis ev + 5 között fogad el "
        "értékeket. Az értékelés 0.0 és 10.0 közötti DECIMAL érték "
        "(0.1 lepeskozzel).")

    add_paragraph_styled(doc,
        "A mentés után a rendszer a crud_create action-on keresztül "
        "szurja be az adatokat a filmek táblába prepared statement "
        "segitsegevel, majd flash üzenettel tájékoztatja a felhasználót "
        "és átirányít a listara.")

    add_screenshot(doc, "5.6.b", "5.6.b. ábra: Uj film hozzáadása űrlap")

    add_heading_styled(doc, "Szerkesztés (Update)", level=3)
    add_paragraph_styled(doc,
        "A Szerkesztés gombra kattintva (action=szerkeszt&id=X) ugyanaz "
        "az űrlap jelenik meg, mint a létrehozasnal, de a mezők a film "
        "aktualis adataival vannak kitoltve. Az id paraméter alapján a "
        "rendszer lekérdezi a filmet az adatbázisbol. Ha a film nem "
        "található, hibaüzenet jelenik meg. A módosítás a crud_update "
        "action-on keresztül történik UPDATE SQL utasitassal.")

    add_screenshot(doc, "5.6.c", "5.6.c. ábra: Film szerkesztése")

    add_heading_styled(doc, "Törlés (Delete)", level=3)
    add_paragraph_styled(doc,
        "A Törlés gombra kattintva (action=torol&id=X) egy megerosito "
        "oldal jelenik meg, amely megjelenítia a film adatait (cim, rendező, "
        "ev, műfaj, értékelés) és megkerdezi a felhasználót, biztosan "
        "torolni kivanja-e. Az 'Igen, törlés' gombra kattintva a rendszer "
        "a crud_delete action-on keresztül torli a filmet a filmek táblabol. "
        "A 'Megsem' gomb visszairanyit a listara.")

    add_screenshot(doc, "5.6.d", "5.6.d. ábra: Film törlésenek megerositese")

    add_page_break(doc)

    add_heading_styled(doc, "6. Biztonság", level=1)

    add_paragraph_styled(doc,
        "Az alkalmazás több retuert biztonsagi megoldast alkalmaz a "
        "leggyakoribb webes tamadasok ellen. Az alábbiakban reszletezzuk "
        "az egyes vedelmeket.")

    add_heading_styled(doc, "Jelszó hasheles (Password Hashing)", level=2)
    add_paragraph_styled(doc,
        "A felhasználók jelszavait soha nem tároljuk nyers szövegként. "
        "A regisztráció során a password_hash() függvény bcrypt "
        "algoritmussalhaseli a jelszót (PASSWORD_DEFAULT). A bejelentkezés "
        "során a password_verify() függvény hasonlitja ossze a megadott "
        "jelszót a tárolt hash-sel. Ez biztosítja, hogy meg adatbázis "
        "kompromittalas esetén sem fejthetok meg a jelszavak.")

    add_code_block(doc, "// Regisztráció:\n$hashelt_jelszó = password_hash($jelszó, PASSWORD_DEFAULT);\n\n// Bejelentkezés:\nif (password_verify($jelszó, $user['jelszó'])) { ... }")

    add_heading_styled(doc, "Prepared Statements (SQL Injection vedelem)", level=2)
    add_paragraph_styled(doc,
        "Minden adatbázis-lekérdezés PDO prepared statement-eket használ "
        "nevesitett paraméterekkel (:param). Ez teljes vedelmet nyujt "
        "az SQL injection tamadasok ellen, mivel a paraméterek értékei "
        "soha nem epulnek be kozvetlenul az SQL utasitasba. A PDO "
        "ATTR_EMULATE_PREPARES beállítás false-ra van állítva, igy "
        "a szerver oldali prepared statement-ek használodnak.")

    add_code_block(doc, "$stmt = $dbh->prepare(\n    'SELECT id FROM felhasználók WHERE felhasználónév = :fnev'\n);\n$stmt->execute([':fnev' => $felhasználónév]);")

    add_heading_styled(doc, "XSS (Cross-Site Scripting) vedelem", level=2)
    add_paragraph_styled(doc,
        "Minden felhasználói bemenet, amely a HTML kimeneten megjelenik, "
        "az htmlspecialchars() függvényen megy keresztül. Ez a függvény "
        "a <, >, &, \" és ' karaktereket HTML entitásokka alakitja, "
        "megakadályozva a rosszindulatu JavaScript kodok futasat a "
        "felhasználók bongeszojeben.")

    add_code_block(doc, "<?= htmlspecialchars($film['cim']) ?>\n<?= htmlspecialchars($flash['üzenet']) ?>")

    add_heading_styled(doc, "MIME típus ellenőrzés (képfeltöltés)", level=2)
    add_paragraph_styled(doc,
        "A képfeltöltés során a rendszer nem csak a kiterjesztest, hanem "
        "a fájl tenyleg MIME típusat is ellenőrzi a mime_content_type() "
        "függvénnyel. Ez megakadályozza, hogy rosszindulatu fájlokat "
        "(pl. PHP szkripteket) tolthessenek fel kép mezzoben. Csak a "
        "következő típusok engedetyezettek: image/jpeg, image/png, "
        "image/gif, image/webp.")

    add_heading_styled(doc, "Session kezeles", level=2)
    add_paragraph_styled(doc,
        "Az alkalmazás PHP session-oket használ a felhasználói állapot "
        "követesere. A bejelentkezés után a felhasználó adatai "
        "(jelszó nélkül!) a $_SESSION['user'] tömbben tárolódnak. "
        "A kijelentkezéskor a session teljesen megsemmisul "
        "($_SESSION = []; session_destroy();). A flash üzenetek szinten "
        "a session-ben tárolódnak, és automatikusan torlodnek az "
        "elso megjelenitest követoen (egyszeri megjelenites).")

    add_heading_styled(doc, "Engedélyezett oldalak whitelist", level=2)
    add_paragraph_styled(doc,
        "Az index.php egy engedélyezett oldalak listat ($engedélyezett_oldalak) "
        "tart fenn. Csak azok az oldalak érhetők el, amelyek ebben a listaban "
        "szerepelnek. Ha a felhasználó nem letezoeleo oldalt kér, flash "
        "hibaüzenet jelenik meg és átirányítas történik a főoldalra. "
        "Ez megakadályozza a tetszoleges fájl betolteset (Path Traversal).")

    add_page_break(doc)

    add_heading_styled(doc, "7. Reszponzív dizájn", level=1)

    add_paragraph_styled(doc,
        "Az alkalmazás teljes mértékben reszponzív, vagyis alkalmazkodik "
        "a különböző képernyőmeretekhez (asztali gepek, tabletek, mobiltelefonok). "
        "A reszponzív megjelenites három fő technikan alapúl: CSS média queries, "
        "Flexbox layout és CSS Grid.")

    add_heading_styled(doc, "Media queries és breakpointok", level=2)

    add_paragraph_styled(doc,
        "Az alkalmazás két fő breakpointot használ, amelyek a tabletes "
        "és mobil nézeteket választjak el az asztali nézettol:")

    bp_data = [
        ["max-width: 768px", "Tablet", "Hamburger menü, 2 oszlopos galéria, kisebb padding"],
        ["max-width: 480px", "Mobil", "1 oszlopos galéria, kisebb betumeret (14px), kompakt layout"],
    ]
    create_styled_table(doc,
        ["Breakpoint", "Eszkoz", "Valtozasok"],
        bp_data,
        col_widths=[4.5, 3, 9.5])

    doc.add_paragraph()

    add_heading_styled(doc, "Hamburger menü mobil nézetben", level=2)

    add_paragraph_styled(doc,
        "768px alatti képernyőn a navigáció rejtett állapotban van "
        "(max-height: 0; overflow: hidden;) és egy három vonalu hamburger "
        "gomb jelenik meg. A gombra kattintva a JavaScript toggle-eli "
        "az 'active' osztalyt, amely megnyitja a navigációt fuggoleges "
        "elrendezésben. A hamburger gomb animalt X-alakra valtozik "
        "(CSS transform). A menüpont kattintasra automatikusan bezarul "
        "a mobil menü.")

    add_code_block(doc,
        "/* Hamburger gomb megjelenítése */\n"
        ".hamburger { display: flex; }\n\n"
        "/* Navigáció elrejtese */\n"
        "nav { max-height: 0; overflow: hidden; transition: max-height 0.4s ease; }\n\n"
        "/* Navigáció megjelenítése aktiv állapotban */\n"
        "nav.active { max-height: 500px; }")

    add_heading_styled(doc, "CSS Grid galéria", level=2)

    add_paragraph_styled(doc,
        "A képgaléria CSS Grid-et használ a reszponzív elrendezéshez. "
        "Asztali gepen 3 oszlopos, tableten 2 oszlopos, mobilon 1 oszlopos "
        "elrendezést alkalmaz. A galéria elemek hover állapotban enyheon "
        "felskalazodasat (scale(1.03)) és arnyek-melyulest mutatnak, ami "
        "kelolemest ad az interakciohoz.")

    add_code_block(doc,
        "/* Asztali: 3 oszlop */\n"
        ".gallery { display: grid; grid-template-columns: repeat(3, 1fr); gap: 1.2rem; }\n\n"
        "/* Tablet (768px): 2 oszlop */\n"
        "@média (max-width: 768px) {\n"
        "    .gallery { grid-template-columns: repeat(2, 1fr); }\n"
        "}\n\n"
        "/* Mobil (480px): 1 oszlop */\n"
        "@média (max-width: 480px) {\n"
        "    .gallery { grid-template-columns: 1fr; }\n"
        "}")

    add_heading_styled(doc, "Egyeb reszponzív megoldasok", level=2)

    responsive_features = [
        "A Flexbox layout biztosítja a header és navigáció rugalmas elrendezéset",
        "A táblázatok .table-responsive wrapper-ben vannak, igy vizszintesen gorgethetokke valnak kis képernyőn",
        "A képek max-width: 100%-ra vannak állítva, igy soha nem lognak ki a szulo elembol",
        "A videók és térképek szelessege 100%-ra állítva, magassaguk állítodik",
        "A betumeret mobilon 14px-re csokken (html font-size), az összes rem érték aranyosan csokken",
        "A gombok és űrlapmezők merete szinten alkalmazkodik a képernyőmerethez",
    ]
    for rf in responsive_features:
        add_bullet_point(doc, rf)

    add_screenshot(doc, "7", "7. ábra: Az alkalmazás mobil nézetben (hamburger menü és 1 oszlopos galéria)")

    add_page_break(doc)

    add_heading_styled(doc, "8. Összefoglalás", level=1)

    add_paragraph_styled(doc,
        "A Filmtár webalkalmazás egy teljes értékű, PHP alapú "
        "filmadatbázis-kezelő rendszer, amely a modern webfejlesztes "
        "legjobb gyakorlatait alkalmazza. A projekt során a következő "
        "fő célkituzeseket valosiotottuk meg:")

    summary_points = [
        "Teljes CRUD funkciónalitas megvalósítása a filmek kezelésere",
        "Felhasználói autentikáció (regisztráció, bejelentkezés, kijelentkezés) biztonkagos session kezelestel",
        "Front-controller tervezési minta alkalmazása központi routing-gal",
        "PRG (Post-Redirect-Get) minta következetes használata",
        "Többretegu biztonsag: jelszóhasheles, prepared statements, XSS vedelem, MIME ellenőrzés",
        "Ketszintu űrlapvalidáció (kliens oldali JavaScript + szerver oldali PHP)",
        "Teljes mértékben reszponzív dizajn három breakpointtal",
        "Interaktiv elemek: lightbox galéria, hamburger menü, flash üzenetek",
        "Letisztuolt, jol szervezett kodstruktura (sablonok, oldalak, statikus erofforrasok szeparalasa)",
    ]
    for sp in summary_points:
        add_bullet_point(doc, sp)

    add_paragraph_styled(doc,
        "Az alkalmazás továbbfejlesztesi lehetőségei koze tartozik a "
        "filmek kéréseese és szurese, lapozas (paginacio), felhasználói "
        "jogosultsagi szintek (admin/user), filmekhez fuzott kommentek, "
        "valamint RESTful API kialakitasa. A jelenlegi verzio stabil "
        "alapot biztosít ezekhez a bovitesekhez.")

    add_page_break(doc)

    add_heading_styled(doc, "9. Munkafelosztás", level=1)

    add_paragraph_styled(doc,
        "A projektet két fős csoportmunkaban keszitettuk el. Az alábbi táblázat "
        "rögzíti, hogy a két hallgato kozul ki melyik feladatreszt valósította meg.")

    munkafelosztas = [
        ["Hallgató",        "Neptun", "Felelősségi terület"],
        ["Gaál Péter",      "GULX05", "Backend, adatbázis, autentikáció, CRUD, deploy"],
        ["Molnár Ádám",     "MFG82Z", "Frontend, reszponzív CSS, JavaScript, multimédia"],
    ]
    create_styled_table(doc, munkafelosztas[0], munkafelosztas[1:], col_widths=[4, 3, 9])
    doc.add_paragraph()

    add_heading_styled(doc, "Gaál Péter (GULX05) - Backend és infrastruktúra", level=2)
    gp_items = [
        "Adatbázis-sema tervezése: g_felhasználók, g_filmek, g_üzenetek, g_képek táblák és külső kulcs kapcsolatok",
        "sql/database.sql és sql/seed.php adatbázis-inicializációs szkriptek",
        "Front-controller tervezési minta megvalósítása (index.php) - routing, POST feldolgozás, PRG minta",
        "Autentikáció: regisztráció, bejelentkezés, kijelentkezés, session-kezelés",
        "CRUD műveletek a filmek táblához (pages/crud.php) - lista, uj, szerkesztés, törlés útvonalakkal",
        "Bejelentkezési és regisztrációs oldalak (pages/belépés.php, pages/regisztráció.php)",
        "Apache .htaccess és config.php (PDO prepared statements) konfiguráció",
        "Internetes tárhelyre valo telepítés - Nethely.hu, FTP feltöltés, MySQL import",
        "Megosztott tárhelyhez g_ tábla-prefix bevezetése az ütközések elkerülésére",
    ]
    for it in gp_items:
        add_bullet_point(doc, it)

    add_heading_styled(doc, "Molnár Ádám (MFG82Z) - Frontend, multimédia, deploy-csomag", level=2)
    ma_items = [
        "HTML5 sablonok: templates/header.php (navigáció, felhasználói info), templates/footer.php",
        "Teljes reszponzív CSS3 stíluslap (css/style.css) - Flexbox, Grid, két breakpoint média query",
        "JavaScript funkciók (js/main.js): hamburger menü, lightbox képgaléria, flash üzenet animáció",
        "Kapcsolati űrlap kliens oldali validációja (js/validation.js) - regex, hossz-ellenőrzés",
        "Tartalmi oldalak: főoldal, képgaléria + feltöltés, kapcsolati űrlap, üzenetek listázása",
        "Multimédia: 5 másodperces saját intró videó, YouTube beágyazás (Saul fia hivatalos előzetes), Google térkép",
        "Nethely deployment csomag és útmutató (NETHELY_DEPLOY.md, config.nethely.php, sql/database_nethely.sql)",
        "Képgaléria és üzenetek SQL-lekérdezések prefix-átírása",
    ]
    for it in ma_items:
        add_bullet_point(doc, it)

    add_heading_styled(doc, "Közös munka", level=2)
    kozos_items = [
        "Téma-választás (magyar filmadatbázis), tervezés, követelmény-elemzés",
        "Kód-áttekintés, hibajavítás, refaktorálás",
        "Dokumentacio összeállítása, képernyőképek készítése",
        "Tesztelés a hostolt alkalmazáson",
    ]
    for it in kozos_items:
        add_bullet_point(doc, it)

    add_page_break(doc)

    add_heading_styled(doc, "10. Beadási és belépési adatok", level=1)

    add_paragraph_styled(doc,
        "Az alábbi adatok szükségesek az alkalmazás elérhetőséghez és ellenőrzésehez. "
        "A jelszavakat csak az oktató szamara, a megoldas javításához adjuk meg.")

    add_heading_styled(doc, "10.1. Internetes elérhetőség", level=2)
    elérhető_rows = [
        ["Weboldal URL",   "http://filmtar.nhely.hu/"],
        ["GitHub repo",    "https://github.com/Peti352/Filmtár-Webprog"],
    ]
    create_styled_table(doc, ["Megnevezés", "Cím"], elérhető_rows, col_widths=[5, 11])
    doc.add_paragraph()

    add_heading_styled(doc, "10.2. Tárhely (FTP) belépési adatok", level=2)
    ftp_rows = [
        ["FTP host",       "ftp.nethely.hu"],
        ["FTP felhasználó","filmtar"],
        ["FTP jelszó",     "Webprog-1!"],
        ["FTP port",       "21"],
    ]
    create_styled_table(doc, ["Megnevezés", "Érték"], ftp_rows, col_widths=[5, 11])
    doc.add_paragraph()

    add_heading_styled(doc, "10.3. MySQL adatbázis belépési adatok", level=2)
    db_rows = [
        ["DB host",        "localhost"],
        ["DB nev",         "filmtar"],
        ["DB felhasználó", "filmtar"],
        ["DB jelszó",      "Webprog-1!"],
        ["phpMyAdmin",     "https://www.nethely.hu/  (Adatbázis menü -> phpMyAdmin)"],
    ]
    create_styled_table(doc, ["Megnevezés", "Érték"], db_rows, col_widths=[5, 11])
    doc.add_paragraph()

    add_heading_styled(doc, "10.4. Teszt-felhasználók", level=2)
    add_paragraph_styled(doc, "A seed-adatokban három teszt-felhasználó van létrehozva, amelyekkel az alkalmazás funkciói azonnal kiprobalhatok:")
    user_rows = [
        ["admin", "admin123",  "Adminisztrátor (alapértelmezett)"],
        ["teszt", "teszt123",  "Teszt-felhasználó"],
        ["user1", "jelszó123", "Normál felhasználó"],
    ]
    create_styled_table(doc, ["Felhasználónév", "Jelszó", "Szerepkor"], user_rows, col_widths=[4, 4, 8])
    doc.add_paragraph()

    add_paragraph_styled(doc,
        "Megjegyzés: Az alkalmazás reszponzív, igy mobil böngészőből is megfelelően "
        "megtekinthető. A bejelentkezés után az 'Üzenetek' menüpont és a CRUD műveletek "
        "is elérhetővé válnak.")

    add_page_break(doc)

    add_heading_styled(doc, "11. Irodalomjegyzék", level=1)

    references = [
        [
            "PHP Official Documentation",
            "https://www.php.net/manual/en/",
            "A PHP nyelv hivatalos dokumentacioja. Használva: PDO, password_hash(), "
            "password_verify(), session_start(), htmlspecialchars(), filter_var(), "
            "mime_content_type(), move_uploaded_file() függvények referenciakent.",
        ],
        [
            "MDN Web Docs - Mozilla Developer Network",
            "https://developer.mozilla.org/",
            "A webes technologiak (HTML, CSS, JavaScript) atfogo referenciaja. "
            "Használva: CSS Grid, Flexbox, Media Queries, addEventListener(), "
            "classList, DOM manipulacio dokumentaciokent.",
        ],
        [
            "W3Schools Online Web Tutorials",
            "https://www.w3schools.com/",
            "Webtechnologiai oktatóanyagok és referenciak. Használva: HTML5 "
            "elemek, CSS selektorok, JavaScript alapok, PHP szintaxis peldakent.",
        ],
        [
            "MySQL 8.0 Reference Manual",
            "https://dev.mysql.com/doc/refman/8.0/en/",
            "A MySQL adatbázis-kezelő rendszer hivatalos dokumentacioja. "
            "Használva: CREATE TABLE, INSERT, UPDATE, DELETE, FOREIGN KEY, "
            "AUTO_INCREMENT szintaxis referenciakent.",
        ],
        [
            "OWASP - Open Web Application Security Project",
            "https://owasp.org/",
            "Webalkalmazás biztonsagi irányelvek és ajánlások. "
            "Használva: SQL Injection, XSS, Password Storage "
            "legjobb gyakorlatok forrasakent.",
        ],
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

        run_num = p.add_run("[{}] ".format(i))
        set_run_font(run_num, size=11, bold=True)

        run_title = p.add_run(ref[0])
        set_run_font(run_title, size=11, bold=True)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.left_indent = Cm(0.7)
        run_url = p2.add_run(ref[1])
        set_run_font(run_url, size=10, italic=True, color=(0, 100, 200))

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(12)
        p3.paragraph_format.left_indent = Cm(0.7)
        run_desc = p3.add_run(ref[2])
        set_run_font(run_desc, size=10)

    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Filmtar_Dokumentacio.docx")
    doc.save(output_path)
    print("Dokumentacio sikérésen generalva: {}".format(output_path))
    print("A fájl megnyithato Microsoft Word-ben és PDF-be mentheto.")


if __name__ == "__main__":
    generate_documentation()
